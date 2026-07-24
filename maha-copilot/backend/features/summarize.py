"""File Summarization for PDF and Word (.docx).

Extract text, then map-reduce: summarize each chunk, then summarize the
summaries into a final brief. Handles long files without exceeding context.
"""
from __future__ import annotations

import io

from pypdf import PdfReader
from docx import Document

from ..config import get_settings
from ..together_client import get_client

_MAP_SYS = """Summarize this section of a document in 3-4 sentences, preserving
concrete facts, figures, dates, names and directives. {lang_line}"""

_REDUCE_SYS = """You are given section summaries of one document. Produce a single
coherent brief for a government officer with: (1) a 3-4 sentence overview, then
(2) 4-8 bullet key points. Preserve figures, dates and directives. {lang_line}"""


def extract_text(filename: str, data: bytes) -> str:
    name = filename.lower()
    if name.endswith(".pdf"):
        reader = PdfReader(io.BytesIO(data))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    if name.endswith(".docx"):
        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(parts)
    # plain text fallback
    try:
        return data.decode("utf-8", errors="ignore")
    except Exception:
        return ""


def _split(text: str, size: int) -> list[str]:
    text = text.strip()
    return [text[i : i + size] for i in range(0, len(text), size)] or [""]


async def summarize_file(filename: str, data: bytes, language: str = "en") -> dict:
    client = get_client()
    text = extract_text(filename, data)
    if not text.strip():
        return {"error": "Could not extract any text from the file."}

    lang_line = (
        "Write in Marathi." if language == "mr" else "Write in English."
    )
    sections = _split(text, 6000)

    # map
    partials = []
    for sec in sections:
        s = await client.chat(
            _MAP_SYS.format(lang_line=lang_line), sec, temperature=0.2, max_tokens=350
        )
        partials.append(s)

    # reduce
    if len(partials) == 1:
        final = await client.chat(
            _REDUCE_SYS.format(lang_line=lang_line), partials[0],
            temperature=0.2, max_tokens=800,
        )
    else:
        joined = "\n\n".join(f"Section {i+1}: {p}" for i, p in enumerate(partials))
        final = await client.chat(
            _REDUCE_SYS.format(lang_line=lang_line), joined,
            temperature=0.2, max_tokens=900,
        )

    return {
        "filename": filename,
        "chars": len(text),
        "sections": len(sections),
        "summary": final,
    }
